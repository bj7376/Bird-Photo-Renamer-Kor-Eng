# 파일 이름: visual_report.py (v2.1 - 썸네일 이미지 사용)
"""
조류 관찰 데이터를 기반으로 HTML/Word 형식의 시각적 리포트를 생성합니다.
"""

from __future__ import annotations

import base64
import io
import os
import re
import shutil
from datetime import datetime
from typing import Dict, List

from PIL import Image

# ---------------------- 유틸리티 ----------------------

def sanitize_filename(name: str) -> str:
    """파일명에 사용할 수 없는 문자 제거"""
    if not isinstance(name, str):
        return ""
    name = name.replace('*', '')
    name = re.sub(r'[\\/:"*?<>|]', '', name).strip()
    return re.sub(r"\s+", "_", name)


def image_to_base64(image_path: str, max_size: tuple = (800, 600)) -> str:
    """이미지를 base64로 인코딩 (HTML 임베딩용)"""
    try:
        with Image.open(image_path) as img:
            # EXIF orientation 처리
            if hasattr(img, '_getexif'):
                exif = img._getexif()
                if exif and 274 in exif:
                    orientation = exif[274]
                    if orientation == 3:
                        img = img.rotate(180, expand=True)
                    elif orientation == 6:
                        img = img.rotate(270, expand=True)
                    elif orientation == 8:
                        img = img.rotate(90, expand=True)
            
            # 리사이즈
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
            
            # base64 인코딩
            buffer = io.BytesIO()
            img_format = 'JPEG' if img.mode == 'RGB' else 'PNG'
            img.save(buffer, format=img_format, quality=85, optimize=True)
            img_data = buffer.getvalue()
            
            mime_type = f"image/{img_format.lower()}"
            return f"data:{mime_type};base64,{base64.b64encode(img_data).decode()}"
    except Exception as e:
        print(f"이미지 base64 변환 실패 ({image_path}): {e}")
        return ""


def get_observation_time_info(observations: List[Dict]) -> Dict[str, str]:
    """관찰 시간 정보 계산 (여러 날 지원)"""
    dates_with_time = [o['datetime'] for o in observations if o['datetime']]
    if not dates_with_time:
        return {
            'date': '관찰 시간 정보 없음',
            'start_time': '',
            'end_time': '',
            'time_range': '',
            'date_range': ''
        }
    
    start_time = min(dates_with_time)
    end_time = max(dates_with_time)
    
    # 같은 날인지 확인
    if start_time.date() == end_time.date():
        # 하루 관찰
        observation_date = start_time.strftime('%Y년 %m월 %d일')
        start_time_str = start_time.strftime('%H:%M')
        end_time_str = end_time.strftime('%H:%M')
        return {
            'date': observation_date,
            'start_time': start_time_str,
            'end_time': end_time_str,
            'time_range': f"{start_time_str} - {end_time_str}",
            'date_range': observation_date
        }
    else:
        # 여러 날 관찰
        start_date_str = start_time.strftime('%Y년 %m월 %d일')
        end_date_str = end_time.strftime('%Y년 %m월 %d일')
        start_time_str = start_time.strftime('%m월 %d일 %H:%M')
        end_time_str = end_time.strftime('%m월 %d일 %H:%M')
        return {
            'date': f"{start_date_str} ~ {end_date_str}",
            'start_time': start_time_str,
            'end_time': end_time_str,
            'time_range': f"{start_time_str} - {end_time_str}",
            'date_range': f"{start_date_str} ~ {end_date_str}"
        }

# --------------------- HTML 리포트 생성 ---------------------

def create_html_report(log_dir: str, observations: List[Dict], location: str, thumbnail_dir: str, thumbnail_size: str, log):
    """HTML 형식의 시각적 리포트 생성"""
    if not observations:
        log("- HTML 리포트를 생성할 기록이 없습니다.")
        return
    
    os.makedirs(log_dir, exist_ok=True)
    
    # 썸네일 크기 설정
    thumb_sizes = {
        'small': (150, 150),
        'medium': (250, 250),
        'large': (400, 400)
    }
    thumb_size_px = thumb_sizes.get(thumbnail_size, (250, 250))
    
    # 관찰 시간 정보
    time_info = get_observation_time_info(observations)
    
    # 종별로 그룹화
    species_groups = {}
    for o in observations:
        key = o['scientific_name']
        if key not in species_groups:
            species_groups[key] = []
        species_groups[key].append(o)
    
    # 분류학적 순서로 정렬
    sorted_species = sorted(species_groups.items(), 
                          key=lambda x: (x[1][0]['taxonomy'].get('order', 'zzz'),
                                       x[1][0]['taxonomy'].get('family', 'zzz')))
    
    # HTML 내용 생성
    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>조류 관찰 보고서</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }}
        .header {{
            text-align: center;
            border-bottom: 3px solid #2c5530;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        .header h1 {{
            color: #2c5530;
            margin: 0;
            font-size: 2.5em;
        }}
        .summary {{
            background: #e8f5e8;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 30px;
        }}
        .summary-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 15px;
        }}
        .summary-item {{
            text-align: center;
        }}
        .summary-number {{
            font-size: 2em;
            font-weight: bold;
            color: #2c5530;
        }}
        .species-section {{
            margin-bottom: 40px;
            border: 1px solid #ddd;
            border-radius: 8px;
            overflow: hidden;
        }}
        .species-header {{
            background: #2c5530;
            color: white;
            padding: 15px 20px;
        }}
        .species-title {{
            margin: 0;
            font-size: 1.4em;
        }}
        .species-info {{
            font-size: 0.9em;
            opacity: 0.9;
            margin-top: 5px;
        }}
        .species-content {{
            padding: 20px;
        }}
        .observation-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(300px, 1fr));
            gap: 20px;
        }}
        .observation-card {{
            border: 1px solid #eee;
            border-radius: 8px;
            overflow: hidden;
            background: #fafafa;
        }}
        .thumb-image {{
            width: 100%;
            height: {thumb_size_px[1]}px;
            object-fit: cover;
            background: #f0f0f0;
        }}
        .observation-info {{
            padding: 15px;
        }}
        .datetime {{
            font-weight: bold;
            color: #2c5530;
            margin-bottom: 10px;
        }}
        .taxonomy {{
            display: grid;
            grid-template-columns: repeat(2, 1fr);
            gap: 10px;
            margin-top: 10px;
            font-size: 0.9em;
        }}
        .taxonomy-item {{
            background: white;
            padding: 8px;
            border-radius: 4px;
            border-left: 3px solid #2c5530;
        }}
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            color: #666;
            font-size: 0.9em;
        }}
        @media print {{
            body {{ background: white; }}
            .container {{ box-shadow: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🐦 조류 관찰 보고서</h1>
            <p>관찰일: {time_info['date']}</p>
            <p>관찰시간: {time_info['time_range']}</p>
            <p>관찰 장소: {location}</p>
        </div>
        
        <div class="summary">
            <h2>📊 관찰 요약</h2>
            <div class="summary-grid">
                <div class="summary-item">
                    <div class="summary-number">{len(observations)}</div>
                    <div>관찰 건수</div>
                </div>
                <div class="summary-item">
                    <div class="summary-number">{len(species_groups)}</div>
                    <div>관찰 종수</div>
                </div>
                <div class="summary-item">
                    <div class="summary-number">{len(set(o['taxonomy'].get('family', 'N/A') for o in observations))}</div>
                    <div>관찰 과수</div>
                </div>
                <div class="summary-item">
                    <div class="summary-number">{len(set(o['taxonomy'].get('order', 'N/A') for o in observations))}</div>
                    <div>관찰 목수</div>
                </div>
            </div>
        </div>
"""
    
    # 각 종별 섹션 생성
    for sci_name, species_observations in sorted_species:
        first_obs = species_observations[0]
        korean_name = first_obs['korean_name']
        common_name = first_obs['common_name']
        order = first_obs['taxonomy'].get('order', 'N/A')
        family = first_obs['taxonomy'].get('family', 'N/A')
        
        html_content += f"""
        <div class="species-section">
            <div class="species-header">
                <h2 class="species-title">{korean_name}</h2>
                <div class="species-info">
                    {common_name} | <em>{sci_name}</em><br>
                    목: {order} | 과: {family}
                </div>
            </div>
            <div class="species-content">
                <div class="observation-grid">
"""
        
        for obs_data in species_observations:
            # 날짜 정보가 여러 날에 걸쳐 있으면 날짜도 함께 표시
            if obs_data['datetime']:
                dates_in_species = [o['datetime'] for o in species_observations if o['datetime']]
                if dates_in_species and min(dates_in_species).date() != max(dates_in_species).date():
                    time_str = obs_data['datetime'].strftime('%m/%d %H:%M:%S')
                else:
                    time_str = obs_data['datetime'].strftime('%H:%M:%S')
            else:
                time_str = '시간 정보 없음'
            
            # 각 관찰 기록의 고유한 썸네일 이미지를 찾습니다.
            thumb_img_path = None
            base_thumb_name = os.path.splitext(obs_data['new_filename'])[0]
            thumb_filename = f"{base_thumb_name}_thumb.jpg"
            potential_path = os.path.join(thumbnail_dir, thumb_filename)
            if os.path.exists(potential_path):
                thumb_img_path = potential_path
            
            # 이미지를 base64로 인코딩
            img_data = ""
            if thumb_img_path:
                img_data = image_to_base64(thumb_img_path, thumb_size_px)
            
            html_content += f"""
                    <div class="observation-card">
                        {f'<img src="{img_data}" alt="{korean_name}" class="thumb-image">' if img_data else '<div class="thumb-image" style="display:flex;align-items:center;justify-content:center;color:#999;">이미지 없음</div>'}
                        <div class="observation-info">
                            <div class="datetime">🕐 {time_str}</div>
                            <div class="taxonomy">
                                <div class="taxonomy-item">
                                    <strong>목:</strong> {order}
                                </div>
                                <div class="taxonomy-item">
                                    <strong>과:</strong> {family}
                                </div>
                            </div>
                        </div>
                    </div>
"""
        
        html_content += """
                </div>
            </div>
        </div>
"""
    
    html_content += """
        <div class="footer">
            <p>본 보고서는 AI 조류 사진 자동 분류 프로그램 v2.1로 생성되었습니다.</p>
            <p>Powered by Google Gemini + Wikipedia</p>
        </div>
    </div>
</body>
</html>
"""
    
    # HTML 파일 저장
    html_path = os.path.join(log_dir, 'visual_report.html')
    try:
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        log(f"  - HTML 리포트 생성 완료: {os.path.basename(html_path)}")
    except Exception as e:
        log(f"  - HTML 리포트 생성 실패: {e}")


# --------------------- Word 리포트 생성 ---------------------

def create_word_report(log_dir: str, observations: List[Dict], location: str, thumbnail_dir: str, log):
    """Word 형식의 시각적 리포트 생성"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import qn
    except ImportError:
        log("  - python-docx 라이브러리가 필요합니다. 'pip install python-docx'로 설치하세요.")
        return
    
    if not observations:
        log("- Word 리포트를 생성할 기록이 없습니다.")
        return
    
    os.makedirs(log_dir, exist_ok=True)
    
    # 새 문서 생성
    doc = Document()
    
    # 한글 폰트 설정 함수 (폴백 지원)
    def set_korean_font(run, font_name=None):
        """한글 텍스트에 적절한 폰트 설정 (폰트 폴백 지원)"""
        font_candidates = [
            "맑은 고딕", "Apple SD Gothic Neo", "Noto Sans CJK KR",
            "Arial Unicode MS", "DejaVu Sans"
        ]
        if font_name:
            font_candidates.insert(0, font_name)
        for font in font_candidates:
            try:
                run.font.name = font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
                break
            except:
                continue
    
    # 관찰 시간 정보
    time_info = get_observation_time_info(observations)
    
    # 문서 제목
    title = doc.add_heading('🐦 조류 관찰 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs: set_korean_font(run)
    
    # 기본 정보
    info_para = doc.add_paragraph()
    run1 = info_para.add_run(f"관찰일: {time_info['date']}\n"); run1.bold = True; set_korean_font(run1)
    run2 = info_para.add_run(f"관찰시간: {time_info['time_range']}\n"); run2.bold = True; set_korean_font(run2)
    run3 = info_para.add_run(f"관찰 장소: {location}"); run3.bold = True; set_korean_font(run3)
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 요약 테이블
    summary_heading = doc.add_heading('📊 관찰 요약', level=1)
    for run in summary_heading.runs: set_korean_font(run)
    
    summary_table = doc.add_table(rows=2, cols=4); summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['관찰 건수', '관찰 종수', '관찰 과수', '관찰 목수']
    values = [
        str(len(observations)),
        str(len(set(o['scientific_name'] for o in observations))),
        str(len(set(o['taxonomy'].get('family', 'N/A') for o in observations))),
        str(len(set(o['taxonomy'].get('order', 'N/A') for o in observations)))
    ]
    
    for i, header in enumerate(headers):
        cell_h = summary_table.cell(0, i); cell_h.text = header
        for p in cell_h.paragraphs:
            for r in p.runs: r.bold = True; set_korean_font(r)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_v = summary_table.cell(1, i); cell_v.text = values[i]
        for p in cell_v.paragraphs:
            for r in p.runs: set_korean_font(r)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 종별 섹션
    species_groups = {}
    for o in observations:
        key = o['scientific_name']; species_groups.setdefault(key, []).append(o)
    
    sorted_species = sorted(species_groups.items(), key=lambda x: (x[1][0]['taxonomy'].get('order', 'zzz'), x[1][0]['taxonomy'].get('family', 'zzz')))
    
    record_heading = doc.add_heading('🔍 종별 관찰 기록', level=1)
    for run in record_heading.runs: set_korean_font(run)
    
    for sci_name, species_observations in sorted_species:
        first_obs = species_observations[0]
        korean_name = first_obs['korean_name']
        common_name = first_obs['common_name']
        order = first_obs['taxonomy'].get('order', 'N/A')
        family = first_obs['taxonomy'].get('family', 'N/A')
        
        species_title = doc.add_heading(f"{korean_name}", level=2)
        for run in species_title.runs: set_korean_font(run)
        
        species_info = doc.add_paragraph()
        run1 = species_info.add_run(f"{common_name} | "); run1.italic = True; set_korean_font(run1)
        run2 = species_info.add_run(f"{sci_name}\n"); run2.italic = True
        run3 = species_info.add_run(f"목: {order} | 과: {family}"); set_korean_font(run3)
        
        table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_texts = ['썸네일 이미지', '관찰 시간', '분류 정보']
        for i, text in enumerate(header_texts):
            cell = header_cells[i]; cell.text = text
            for p in cell.paragraphs:
                for r in p.runs: r.bold = True; set_korean_font(r)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for obs_data in species_observations:
            row_cells = table.add_row().cells
            
            # 각 관찰 기록의 고유한 썸네일 이미지를 찾습니다.
            thumb_img_path = None
            base_thumb_name = os.path.splitext(obs_data['new_filename'])[0]
            thumb_filename = f"{base_thumb_name}_thumb.jpg"
            potential_path = os.path.join(thumbnail_dir, thumb_filename)
            if os.path.exists(potential_path):
                thumb_img_path = potential_path
            
            if thumb_img_path:
                try:
                    p = row_cells[0].paragraphs[0]
                    r = p.runs[0] if p.runs else p.add_run()
                    r.add_picture(thumb_img_path, width=Inches(1.5))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    row_cells[0].text = "이미지 로드 실패"; log(f"  - Word 이미지 삽입 실패: {e}")
            else:
                row_cells[0].text = "이미지 없음"
            
            # 모든 텍스트 셀에 폰트 적용
            for p in row_cells[0].paragraphs:
                for r in p.runs: set_korean_font(r)

            if obs_data['datetime']:
                dates_in_species = [o['datetime'] for o in species_observations if o['datetime']]
                if dates_in_species and min(dates_in_species).date() != max(dates_in_species).date():
                    time_str = obs_data['datetime'].strftime('%m/%d %H:%M:%S')
                else:
                    time_str = obs_data['datetime'].strftime('%H:%M:%S')
            else:
                time_str = '시간 정보 없음'
            
            row_cells[1].text = time_str;
            for p in row_cells[1].paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER; [set_korean_font(r) for r in p.runs]
            
            row_cells[2].text = f"목: {order}\n과: {family}"
            for p in row_cells[2].paragraphs: [set_korean_font(r) for r in p.runs]

        doc.add_paragraph()
    
    word_path = os.path.join(log_dir, 'visual_report.docx')
    try:
        doc.save(word_path)
        log(f"  - Word 리포트 생성 완료: {os.path.basename(word_path)}")
    except Exception as e:
        log(f"  - Word 리포트 생성 실패: {e}")


# --------------------- 메인 인터페이스 ---------------------

def create_visual_reports(observations: List[Dict], out_dir: str, src_dir: str, report_options: Dict, location: str, log):
    """시각적 리포트 생성 메인 함수"""
    log_dir = os.path.join(out_dir, '탐조기록')
    thumbnail_dir = os.path.join(out_dir, 'thumbnail_images')
    
    report_format = report_options.get('format', 'html')
    thumbnail_size = report_options.get('thumbnail_size', 'medium')
    
    if report_format in ['html', 'both']:
        log("- HTML 시각적 리포트 생성 중...")
        create_html_report(log_dir, observations, location, thumbnail_dir, thumbnail_size, log)
    
    if report_format in ['docx', 'both']:
        log("- Word 시각적 리포트 생성 중...")
        create_word_report(log_dir, observations, location, thumbnail_dir, log)